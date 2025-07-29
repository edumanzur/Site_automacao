from docx import Document

def substituir_em_runs_preservando_tudo(paragrafos, dados):
    for paragrafo in paragrafos:
        texto_acumulado = ""
        run_indices = []
        i = 0

        while i < len(paragrafo.runs):
            run = paragrafo.runs[i]
            texto_acumulado += run.text
            run_indices.append(i)

            for chave, valor in dados.items():
                placeholder = f'{{{{{chave}}}}}'
                if placeholder in texto_acumulado:
                    # Encontrou o placeholder
                    inicio = texto_acumulado.index(placeholder)
                    fim = inicio + len(placeholder)

                    # Reconstrói o texto com o valor no lugar do placeholder
                    novo_texto = texto_acumulado[:inicio] + valor + texto_acumulado[fim:]

                    # Atualiza os runs envolvidos
                    texto_restante = novo_texto
                    for idx in run_indices:
                        paragrafo.runs[idx].text = texto_restante[:len(paragrafo.runs[idx].text)]
                        texto_restante = texto_restante[len(paragrafo.runs[idx].text):]

                    # Se sobrar texto, adiciona em um novo run
                    if texto_restante:
                        paragrafo.runs[run_indices[-1]].add_text(texto_restante)

                    # Reset para continuar o loop
                    texto_acumulado = ""
                    run_indices = []
                    i = -1  # reinicia do começo
                    break  # sai do loop de chaves
            i += 1


def preencher_modelo(caminho_modelo, caminho_saida, dados):
    doc = Document(caminho_modelo)

    substituir_em_runs_preservando_tudo(doc.paragraphs, dados)

    for tabela in doc.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                substituir_em_runs_preservando_tudo(celula.paragraphs, dados)

    for section in doc.sections:
        for parte in [section.header, section.footer]:
            substituir_em_runs_preservando_tudo(parte.paragraphs, dados)

    doc.save(caminho_saida)
    print(f"Arquivo gerado com sucesso em: {caminho_saida}")

