from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.responses import Response
from fastapi.middleware.cors import CORSMiddleware
import os
import tempfile
import shutil
from pathlib import Path

# Importar suas funções
try:
    from extracao import extrair_campos
    print("✅ Módulo extracao importado com sucesso")
    print("⚠️ ATENÇÃO: extrair_campos é uma função ASYNC")
except ImportError as e:
    print(f"❌ Erro ao importar extracao: {e}")
    async def extrair_campos(pdf_path):
        return {"erro": "Módulo extracao não encontrado"}

try:
    from preencher import preencher_modelo
    print("✅ Módulo preencher importado com sucesso")
    print("📋 preencher_modelo precisa de: (caminho_modelo, caminho_saida, dados)")
except ImportError as e:
    print(f"❌ Erro ao importar preencher: {e}")
    def preencher_modelo(*args, **kwargs):
        raise Exception("Módulo preencher não encontrado")

app = FastAPI()

# CORS
origins = [
    "http://localhost:3000",  # para dev local
    "https://site-automacao-r80e36cht-edumanzur21-6375s-projects.vercel.app",
    "htpps://site-automacao-eight.vercel.app"
]
app.add_middleware(
    CORSMiddleware,
    allow_origins=origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

@app.get("/")
async def root():
    return {"message": "Servidor funcionando!"}

@app.get("/test-docx")
async def test_docx():
    """Rota de teste para verificar se a geração de DOCX funciona"""
    try:
        # Criar um documento simples na memória
        doc = Document()
        doc.add_heading('Documento de Teste', 0)
        doc.add_paragraph('Este é um teste de geração de DOCX.')
        doc.add_paragraph('Se você está lendo isso, a geração funcionou!')
        
        # Salvar em buffer de memória
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        # Retornar como resposta direta
        return Response(
            content=buffer.getvalue(),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": "attachment; filename=teste.docx"
            }
        )
    except Exception as e:
        print(f"Erro no teste: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/upload-pdf")
async def upload_pdf(file: UploadFile = File(...)):
    print(f"=== INÍCIO DO PROCESSAMENTO ===")
    print(f"Arquivo recebido: {file.filename}")
    print(f"Content-type: {file.content_type}")
    
    # Verificar se é PDF
    if not file.filename.lower().endswith('.pdf'):
        raise HTTPException(status_code=400, detail="Apenas arquivos PDF são aceitos")
    
    # Criar diretório temporário
    temp_dir = tempfile.mkdtemp()
    input_path = None
    output_path = None
    
    try:
        # Salvar arquivo PDF temporariamente
        input_path = os.path.join(temp_dir, f"input_{file.filename}")
        print(f"Salvando PDF em: {input_path}")
        
        with open(input_path, "wb") as buffer:
            content = await file.read()
            buffer.write(content)
        
        print(f"PDF salvo: {len(content)} bytes")
        
        # Verificar se arquivo foi salvo
        if not os.path.exists(input_path):
            raise Exception("Falha ao salvar o arquivo PDF")
        
        # === USAR SUAS FUNÇÕES ===
        print("🔍 Iniciando extração com sua função...")
        try:
            print(f"Chamando: await extrair_campos('{input_path}')")
            # SUA FUNÇÃO É ASYNC - precisa do await
            campos_extraidos = await extrair_campos(input_path)
            
            print(f"✅ Extração concluída: {type(campos_extraidos)}")
            
            if isinstance(campos_extraidos, dict):
                print(f"📋 Campos extraídos: {list(campos_extraidos.keys())}")
                print(f"📊 Total de campos: {len(campos_extraidos)}")
                
                # Mostrar amostra dos dados
                for chave, valor in list(campos_extraidos.items())[:5]:  # Primeiros 5 campos
                    valor_str = str(valor)[:100] + "..." if len(str(valor)) > 100 else str(valor)
                    print(f"   {chave}: {valor_str}")
                
                if len(campos_extraidos) > 5:
                    print(f"   ... e mais {len(campos_extraidos) - 5} campos")
                    
            else:
                print(f"⚠️ Tipo inesperado retornado: {type(campos_extraidos)}")
                # Normalizar para dict se necessário
                campos_extraidos = {
                    "dados_extraidos": str(campos_extraidos),
                    "tipo_original": str(type(campos_extraidos)),
                    "arquivo_fonte": file.filename
                }
                
        except Exception as e:
            print(f"❌ Erro na extração: {e}")
            import traceback
            print(f"Traceback da extração: {traceback.format_exc()}")
            
            # Dados de fallback baseados na sua estrutura
            campos_extraidos = {
                "PACIENTE": "Erro na extração",
                "NASCIMENTO": "Não encontrado", 
                "IDADE": "Não encontrado",
                "MAE": "Não encontrado",
                "TELEFONE": "Não encontrado",
                "CRM": "Não encontrado",
                "MEDICO": "Não encontrado",
                "DIAGNOSTICO": "Não encontrado",
                "PROCEDIMENTO": "Não encontrado",
                "NACIONALIDADE": "Não encontrado",
                "CEP": "Não encontrado",
                "RISCO": "Não encontrado",
                "DIAGNOSTICO_INICIAL": "Não encontrado",
                "CENTRAL": "Não encontrado",
                "UNIDADE": "Não encontrado",
                "DATA_SOLICITACAO": "Não encontrado",
                "SITUACAO_ATUAL": "Não encontrado",
                "HOSPITAL": "Não encontrado",
                "CODIGO_SOLICITACAO": "Não encontrado",
                "erro_extracao": str(e),
                "arquivo_fonte": file.filename
            }
        
        # Definir caminho de saída
        output_filename = f"convertido_{file.filename.replace('.pdf', '.docx')}"
        output_path = os.path.join(temp_dir, output_filename)
        
        print(f"📝 Iniciando preenchimento com sua função...")
        print(f"Output path: {output_path}")
        
        try:
            # Sua função precisa de: preencher_modelo(caminho_modelo, caminho_saida, dados)
            
            # Verificar se existe um template
            template_path = "template.docx"  # Nome padrão
            possible_templates = [
                "template.docx",
                "modelo.docx", 
                "template/template.docx",
                "templates/template.docx"
            ]
            
            template_encontrado = None
            for possible_template in possible_templates:
                if os.path.exists(possible_template):
                    template_encontrado = possible_template
                    break
            
            if template_encontrado:
                print(f"📄 Template encontrado: {template_encontrado}")
                
                # Chamar sua função com os parâmetros corretos
                preencher_modelo(template_encontrado, output_path, campos_extraidos)
                print("✅ Função preencher_modelo executada com sucesso")
                
            else:
                print("❌ Nenhum template encontrado!")
                print(f"Procurado em: {possible_templates}")
                raise Exception(f"Template não encontrado. Coloque um arquivo template.docx no diretório do projeto.")
                
        except Exception as e:
            print(f"❌ Erro no preenchimento: {e}")
            print(f"Tentando método alternativo...")
            
            # Fallback: criar documento simples
            try:
                from docx import Document
                doc = Document()
                doc.add_heading('Conversão de PDF', 0)
                doc.add_paragraph(f'Arquivo: {file.filename}')
                doc.add_paragraph('Erro na conversão personalizada.')
                doc.add_paragraph(f'Erro: {str(e)}')
                
                # Adicionar dados extraídos de forma segura
                doc.add_heading('Dados Extraídos:', level=1)
                if isinstance(campos_extraidos, dict):
                    for k, v in campos_extraidos.items():
                        try:
                            doc.add_paragraph(f'{k}: {str(v)[:500]}')
                        except Exception:
                            doc.add_paragraph(f'{k}: [Erro ao converter valor]')
                elif isinstance(campos_extraidos, str):
                    doc.add_paragraph(f'Texto extraído: {campos_extraidos[:1000]}')
                else:
                    doc.add_paragraph(f'Dados: {str(campos_extraidos)[:500]}')
                    
                doc.save(output_path)
                print("📄 Documento fallback criado")
            except Exception as e2:
                raise Exception(f"Erro crítico: {e}. Fallback também falhou: {e2}")
        
        # Verificar se arquivo foi criado
        if not os.path.exists(output_path):
            raise Exception("Arquivo DOCX não foi gerado")
        
        file_size = os.path.getsize(output_path)
        print(f"✅ DOCX criado: {output_path} ({file_size} bytes)")
        
        # Ler arquivo para retornar
        with open(output_path, "rb") as f:
            docx_content = f.read()
        
        print(f"📤 Enviando arquivo: {len(docx_content)} bytes")
        
        return Response(
            content=docx_content,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f"attachment; filename={file.filename.replace('.pdf', '')}_convertido.docx",
                "Content-Length": str(len(docx_content)),
                "Cache-Control": "no-cache"
            }
        )
        
    except Exception as e:
        print(f"❌ ERRO CRÍTICO: {e}")
        import traceback
        print(f"Traceback: {traceback.format_exc()}")
        raise HTTPException(status_code=500, detail=f"Erro na conversão: {str(e)}")
    
    finally:
        # Limpar arquivos temporários
        try:
            if temp_dir and os.path.exists(temp_dir):
                shutil.rmtree(temp_dir)
                print(f"🧹 Diretório temporário removido: {temp_dir}")
        except Exception as e:
            print(f"⚠️ Erro ao limpar arquivos temporários: {e}")

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
